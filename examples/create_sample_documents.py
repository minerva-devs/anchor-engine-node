# Sample Documents Creation Script

import os
import tempfile
from docx import Document

# Create sample_documents directory if it doesn't exist
os.makedirs("sample_documents", exist_ok=True)

# Create a sample text file
with open("sample_documents/project_report.txt", "w") as f:
    f.write("""
Project Status Report
=====================

Project Manager: John Smith
Project Name: AI Implementation
Start Date: 01/15/2024
Deadline: 12/31/2024

Executive Summary
-----------------
The AI Implementation project is progressing according to plan. We have successfully completed the initial research phase and 
have moved into the development phase. Key team members include Jane Doe (Lead Developer) and Bob Johnson (Data Scientist).

Current Status
--------------
- Research Phase: Completed (01/15/2024 - 03/15/2024)
- Development Phase: In Progress (03/16/2024 - 09/30/2024)
- Testing Phase: Pending (10/01/2024 - 11/30/2024)
- Deployment: Pending (12/01/2024 - 12/31/2024)

Budget
------
Total Budget: $500,000
Spent to Date: $275,000
Remaining: $225,000

Key Milestones
--------------
1. Initial Research Completion: 03/15/2024
2. Prototype Development: 06/30/2024
3. System Testing: 11/15/2024
4. Final Deployment: 12/31/2024

Contact Information
-------------------
Project Manager: John Smith (john.smith@company.com)
Lead Developer: Jane Doe (jane.doe@company.com)
Data Scientist: Bob Johnson (bob.johnson@company.com)
Phone: 555-123-4567

Next Review Meeting: 04/15/2024
""")

# Create a sample DOCX file
doc = Document()
doc.add_heading('Business Plan', 0)

doc.add_heading('Executive Summary', level=1)
doc.add_paragraph('Our company is projected to achieve $2.5 million in revenue in the first year, with a 25% profit margin.')
doc.add_paragraph('We plan to capture 5% of the local market within 18 months.')

doc.add_heading('Market Analysis', level=1)
doc.add_paragraph('The target market shows strong growth potential with a 15% annual growth rate.')
doc.add_paragraph('Our competitive advantage lies in our proprietary technology and experienced team.')

doc.add_heading('Financial Projections', level=1)
doc.add_paragraph('Year 1 Revenue: $2.5 million')
doc.add_paragraph('Year 2 Revenue: $5.0 million')
doc.add_paragraph('Year 3 Revenue: $10.0 million')

doc.add_heading('Growth Strategy', level=1)
doc.add_paragraph('Our growth strategy focuses on expanding to three new markets within the first two years.')
doc.add_paragraph('We will invest 20% of revenue in research and development to maintain our technology lead.')

doc.save("sample_documents/business_plan.docx")

# Create a sample HTML file
html_content = """
<!DOCTYPE html>
<html>
<head>
    <title>Meeting Notes</title>
</head>
<body>
    <h1>Project Kickoff Meeting Notes</h1>
    
    <h2>Attendees</h2>
    <ul>
        <li>John Smith - Project Manager</li>
        <li>Jane Doe - Lead Developer</li>
        <li>Bob Johnson - Data Scientist</li>
        <li>Alice Williams - Product Owner</li>
        <li>Charlie Brown - QA Lead</li>
    </ul>
    
    <h2>Agenda Items</h2>
    <ol>
        <li>Project Overview</li>
        <li>Team Introductions</li>
        <li>Timeline Discussion</li>
        <li>Resource Allocation</li>
        <li>Next Steps</li>
    </ol>
    
    <h2>Key Decisions</h2>
    <p>The team agreed on a 12-month timeline with quarterly milestones.</p>
    <p>Jane Doe will lead the development team with Bob Johnson supporting on data science tasks.</p>
    <p>Alice Williams will be the primary contact for product-related questions.</p>
    
    <h2>Action Items</h2>
    <ul>
        <li>John Smith: Create detailed project plan by 01/22/2024</li>
        <li>Jane Doe: Set up development environment by 01/25/2024</li>
        <li>Bob Johnson: Begin data collection by 01/30/2024</li>
        <li>Alice Williams: Define user stories by 02/05/2024</li>
    </ul>
    
    <h2>Next Meeting</h2>
    <p>Date: 02/05/2024</p>
    <p>Location: Conference Room 3</p>
</body>
</html>
"""

with open("sample_documents/meeting_notes.html", "w") as f:
    f.write(html_content)

print("Sample documents created successfully in the 'sample_documents' directory.")